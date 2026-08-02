from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT = Path(r"C:\Users\Hopef\Desktop\carsales\feedback_videos\Client_Feedback_Transcription_and_Requested_Changes.docx")

NAVY = "0B2545"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
GRAY = "555555"
MUTED = "6B7280"
LIGHT_GRAY = "F2F4F7"
LIGHT_BLUE = "E8EEF5"
CALLOUT = "F4F6F9"
BORDER = "D7DBE2"
WHITE = "FFFFFF"


VIDEOS = [
    {
        "number": 1,
        "title": "Vehicle search, sorting, and brand recognition",
        "file": "WhatsApp Video 2026-08-01 at 10.46.28 PM.mp4",
        "duration": "00:37",
        "arabic": (
            "مساء الخير، دكتور، كيف الحال؟ دكتور، هنا في بحث السيارات، شوف يطلع صيني، "
            "استرالي، صيني، بعدين الصين. يعني الصين ما هو مرتب. كذلك بريطانيا تطلع لك "
            "أحياناً ما هي مرتبة، وأمريكا تطلع لك ما هي مرتبة. أنتوا سويتوها بناء على "
            "الحروف ولا بناء على إيش؟ ولو كان فيها لوغو، علامات، يكون أفضل كذا، أفضل من "
            "البحث هذا. وأصلاً لازم فيه أيقونة بحث، بحث عن السيارة على الأقل. كذا الواحد "
            "بيتعب فيها وما راح يلقاها، لأن أصلاً هذه تصير كثيرة مرة."
        ),
        "translation": (
            "Good evening, doctor. How are you? On the vehicle search, look at how the entries "
            "appear: Chinese, Australian, Chinese, then China. The list is not ordered. Britain "
            "sometimes appears out of order, and America also appears without a clear order. "
            "Did you arrange it alphabetically, or by what rule? It would be better if the list "
            "included brand logos or recognizable marks rather than relying only on this search. "
            "There also needs to be a search icon or search field for the vehicle. Otherwise, "
            "people will struggle to find what they need once the list becomes very large."
        ),
        "changes": [
            "Add a clear search icon and search field so users can find a vehicle or brand quickly.",
            "Apply a consistent ordering rule to the country/brand list, such as alphabetical order or a defined business order.",
            "Use brand logos or recognizable marks where possible to make the list easier to scan."
        ],
    },
    {
        "number": 2,
        "title": "Agency/dealer pricing and optional auction price",
        "file": "WhatsApp Video 2026-08-01 at 10.46.29 PM (1).mp4",
        "duration": "00:16",
        "arabic": (
            "طيب بالنسبة لسعر الوكالة، كويس عشان الوكيل يحط سعره أو التاجر. هذي تدل على "
            "سعر المزاد الاختياري، ليش حطيتها؟"
        ),
        "translation": (
            "Regarding the agency price, it is good if the agent or dealer can enter their own "
            "price. But this appears to be an optional auction price. Why was it added?"
        ),
        "changes": [
            "Clarify which user is responsible for entering the agency/dealer price and label the field accordingly.",
            "Confirm the purpose of the optional auction price field; remove it if it is redundant or not part of the intended pricing flow."
        ],
    },
    {
        "number": 3,
        "title": "Custom vehicle categories and colors",
        "file": "WhatsApp Video 2026-08-01 at 10.46.29 PM.mp4",
        "duration": "00:29",
        "arabic": (
            "كذلك دكتور، في صفحة تحكم التاجر، الفئة هنا فيها أشياء محدودة، يعني بالضبط. "
            "يعني يفترض أنه فيه شيء أنا أحطه مانيولي، لأنه فيه أحياناً أسماء غير الأسماء هذه. "
            "بالنسبة للألوان بعد، الألوان لازم فيها بعد مانيولي أنا أتحكم فيها، لأنه فيه أحياناً "
            "ألوان مثلاً تجي مدموجة، لونين مع بعض."
        ),
        "translation": (
            "Also, doctor, on the dealer control page, the trim/category field has only a limited "
            "set of options. There should be a way for me to enter a value manually, because "
            "sometimes the name is different from these options. The same applies to colors: "
            "I should be able to control them manually because a vehicle may have a combined "
            "color - two colors together."
        ),
        "changes": [
            "Allow a dealer to enter a custom/manual trim or category value when the predefined options do not apply.",
            "Allow a custom/manual color value and support combined or two-tone colors."
        ],
    },
    {
        "number": 4,
        "title": "Saudi-market trim naming",
        "file": "WhatsApp Video 2026-08-01 at 10.46.30 PM (1).mp4",
        "duration": "00:13",
        "arabic": (
            "وفي نوع الفئة المتعارف عليه بالسوق السعودي، مش الأسماء هذي، هذي ما يعرفونها. "
            "المتعارف عليه: ستاندر، وفل، ونصف فل."
        ),
        "translation": (
            "There is also a trim classification commonly used in the Saudi market. It is not "
            "these names; people do not recognize those names. The familiar terms are Standard, "
            "Full, and Half Full."
        ),
        "changes": [
            "Replace or supplement the current trim names with the Saudi-market labels: Standard, Full, and Half Full.",
            "Keep the Arabic and English labels consistent wherever trim is displayed or selected."
        ],
    },
    {
        "number": 5,
        "title": "Color checklist and multi-select behavior",
        "file": "WhatsApp Video 2026-08-01 at 10.46.30 PM.mp4",
        "duration": "00:16",
        "arabic": (
            "حتى بالنسبة للألوان، يا دكتور، لو فيه تشيك لست عشان نحدد كمية من الأسود، "
            "من الفضي، من الرمادي، أفضل ما تكون كذا كلها اختيارية."
        ),
        "translation": (
            "For colors as well, doctor, it would be better to have a checklist so we can select "
            "the applicable colors - black, silver, and gray - instead of leaving everything as optional."
        ),
        "changes": [
            "Provide a checklist or multi-select control for common colors such as black, silver, and gray.",
            "Let the dealer select the applicable colors instead of presenting every color as an optional/free-form value."
        ],
    },
    {
        "number": 6,
        "title": "Dealer inventory editing and deletion",
        "file": "WhatsApp Video 2026-08-01 at 10.46.31 PM.mp4",
        "duration": "00:30",
        "arabic": (
            "دكتور، في صفحة التاجر، لما أروح للمخزون مثلاً أبغى أضيف سيارة. أبغى أعدل، "
            "مثلاً أنا أضفت [كلمة غير واضحة في التسجيل]. تمام، أبغى أعدل على الصفحة أو أحذف، "
            "أعدل شيء، أضيف شيء، كيف الطريقة؟ التاجر وارد أنه يخطئ في كلمة، وارد أنه يخطئ "
            "في شيء، فلازم يعدل عليها، وممكن خلاص يبي يحذف الصفحة."
        ),
        "translation": (
            "Doctor, on the dealer page, when I go to inventory, for example, I want to add a car. "
            "If I need to edit it after adding the listing, I should be able to update the page, "
            "delete it, change something, or add something. How is that done? A dealer may make "
            "a typo or make a mistake in some detail, so they need to be able to edit it. They may "
            "also want to delete the listing."
        ),
        "changes": [
            "Add an Edit action for inventory listings so dealers can correct typos and update vehicle details after creation.",
            "Allow dealers to add missing details to an existing listing.",
            "Add a Delete action, ideally with a confirmation step, for listings that should be removed.",
            "Make the inventory edit/delete workflow clear and easy to discover on the dealer page."
        ],
    },
]


SUMMARY_ROWS = [
    ("1", "Search and sorting", "Add vehicle/brand search, apply a consistent list order, and improve scanability with brand logos or marks.", "High"),
    ("2", "Pricing", "Clarify the agency/dealer price field and validate whether the optional auction price field is needed.", "Medium"),
    ("3", "Trim/category", "Allow custom/manual trim or category values when the predefined list does not cover the vehicle.", "High"),
    ("4", "Saudi labels", "Use the Saudi-market trim names Standard, Full, and Half Full, with consistent Arabic/English labels.", "High"),
    ("5", "Colors", "Support a checklist or multi-select for common colors, plus custom and combined/two-tone colors.", "High"),
    ("6", "Inventory actions", "Allow dealers to edit, add details to, and delete inventory listings after creation.", "High"),
]


def rgb(hex_color: str) -> RGBColor:
    return RGBColor.from_string(hex_color)


def ensure_rfonts(parent):
    rpr = parent.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    return rfonts


def set_run_font(run, name="Calibri", size=11, color="000000", bold=None, italic=None):
    run.font.name = name
    rfonts = ensure_rfonts(run._element)
    rfonts.set(qn("w:ascii"), name)
    rfonts.set(qn("w:hAnsi"), name)
    rfonts.set(qn("w:eastAsia"), name)
    rfonts.set(qn("w:cs"), name)
    run.font.size = Pt(size)
    run.font.color.rgb = rgb(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_style_font(style, name="Calibri", size=11, color="000000", bold=None, italic=None):
    style.font.name = name
    style._element.get_or_add_rPr()
    rfonts = ensure_rfonts(style._element)
    rfonts.set(qn("w:ascii"), name)
    rfonts.set(qn("w:hAnsi"), name)
    rfonts.set(qn("w:eastAsia"), name)
    rfonts.set(qn("w:cs"), name)
    style.font.size = Pt(size)
    style.font.color.rgb = rgb(color)
    if bold is not None:
        style.font.bold = bold
    if italic is not None:
        style.font.italic = italic


def set_bidi(paragraph, enabled=True):
    ppr = paragraph._p.get_or_add_pPr()
    bidi = ppr.find(qn("w:bidi"))
    if enabled and bidi is None:
        ppr.append(OxmlElement("w:bidi"))
    elif not enabled and bidi is not None:
        ppr.remove(bidi)


def set_paragraph_border(paragraph, side="left", color=BORDER, size="18", space="8"):
    ppr = paragraph._p.get_or_add_pPr()
    pbdr = ppr.find(qn("w:pBdr"))
    if pbdr is None:
        pbdr = OxmlElement("w:pBdr")
        ppr.append(pbdr)
    border = OxmlElement(f"w:{side}")
    border.set(qn("w:val"), "single")
    border.set(qn("w:sz"), size)
    border.set(qn("w:space"), space)
    border.set(qn("w:color"), color)
    pbdr.append(border)


def set_paragraph_shading(paragraph, fill):
    ppr = paragraph._p.get_or_add_pPr()
    shd = ppr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        ppr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_shading(cell, fill):
    tcpr = cell._tc.get_or_add_tcPr()
    shd = tcpr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tcpr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tcpr = cell._tc.get_or_add_tcPr()
    tc_mar = tcpr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tcpr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color=BORDER, size="6"):
    tblpr = table._tbl.tblPr
    borders = tblpr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tblpr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = qn(f"w:{edge}")
        node = borders.find(tag)
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), size)
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), color)


def set_table_geometry(table, widths, indent=120):
    total = sum(widths)
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl = table._tbl
    tblpr = tbl.tblPr
    tblw = tblpr.find(qn("w:tblW"))
    if tblw is None:
        tblw = OxmlElement("w:tblW")
        tblpr.insert(0, tblw)
    tblw.set(qn("w:w"), str(total))
    tblw.set(qn("w:type"), "dxa")
    tblind = tblpr.find(qn("w:tblInd"))
    if tblind is None:
        tblind = OxmlElement("w:tblInd")
        tblpr.append(tblind)
    tblind.set(qn("w:w"), str(indent))
    tblind.set(qn("w:type"), "dxa")
    layout = tblpr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tblpr.append(layout)
    layout.set(qn("w:type"), "fixed")
    grid = tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            tcpr = cell._tc.get_or_add_tcPr()
            tcw = tcpr.find(qn("w:tcW"))
            if tcw is None:
                tcw = OxmlElement("w:tcW")
                tcpr.insert(0, tcw)
            tcw.set(qn("w:w"), str(widths[idx]))
            tcw.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def keep_row_together(row):
    trpr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    trpr.append(cant_split)


def mark_header_row(row):
    trpr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    trpr.append(header)


def add_page_field(paragraph, field_code):
    run = paragraph.add_run()
    set_run_font(run, size=9, color=MUTED)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = field_code
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.append(begin)
    run._r.append(instr)
    run._r.append(separate)
    run._r.append(text)
    run._r.append(end)


def add_numbering_definition(doc):
    numbering = doc.part.numbering_part.element
    abstract_id = 98
    num_id = 98
    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    level = OxmlElement("w:lvl")
    level.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    numfmt = OxmlElement("w:numFmt")
    numfmt.set(qn("w:val"), "bullet")
    lvltext = OxmlElement("w:lvlText")
    lvltext.set(qn("w:val"), "•")
    lvljust = OxmlElement("w:lvlJc")
    lvljust.set(qn("w:val"), "left")
    ppr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "720")
    tabs.append(tab)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "720")
    ind.set(qn("w:hanging"), "360")
    ppr.append(tabs)
    ppr.append(ind)
    rpr = OxmlElement("w:rPr")
    rfonts = OxmlElement("w:rFonts")
    rfonts.set(qn("w:ascii"), "Symbol")
    rfonts.set(qn("w:hAnsi"), "Symbol")
    rpr.append(rfonts)
    for node in (start, numfmt, lvltext, lvljust, ppr, rpr):
        level.append(node)
    abstract.append(level)
    numbering.append(abstract)
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abs_ref = OxmlElement("w:abstractNumId")
    abs_ref.set(qn("w:val"), str(abstract_id))
    num.append(abs_ref)
    numbering.append(num)
    return num_id


def apply_bullet(paragraph, num_id):
    ppr = paragraph._p.get_or_add_pPr()
    numpr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    numid = OxmlElement("w:numId")
    numid.set(qn("w:val"), str(num_id))
    numpr.append(ilvl)
    numpr.append(numid)
    ppr.append(numpr)


def add_bullet(doc, text, num_id, size=11):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.167
    apply_bullet(p, num_id)
    run = p.add_run(text)
    set_run_font(run, size=size)
    return p


def add_label_paragraph(doc, label, text, label_color=BLUE, text_color=GRAY, size=10):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.10
    r1 = p.add_run(label)
    set_run_font(r1, size=size, color=label_color, bold=True)
    r2 = p.add_run(text)
    set_run_font(r2, size=size, color=text_color)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.paragraph_format.keep_with_next = True
    r = p.add_run(text)
    if level == 1:
        set_run_font(r, size=16, color=BLUE, bold=True)
    elif level == 2:
        set_run_font(r, size=13, color=BLUE, bold=True)
    else:
        set_run_font(r, size=12, color=DARK_BLUE, bold=True)
    return p


def add_table_text(cell, text, size=10, color="000000", bold=False, align=WD_ALIGN_PARAGRAPH.LEFT):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.05
    r = p.add_run(text)
    set_run_font(r, size=size, color=color, bold=bold)
    return p


def build_metadata_table(doc):
    table = doc.add_table(rows=2, cols=4)
    set_table_geometry(table, [1320, 3360, 1320, 3360])
    set_table_borders(table, BORDER, "6")
    # Treat the first metadata row as the repeating header for accessibility tools.
    mark_header_row(table.rows[0])
    rows = [
        ("Project", "Car Sales Platform", "Source", "6 Arabic feedback videos"),
        ("Total audio", "02:22", "Reviewed", "02 Aug 2026"),
    ]
    for row, values in zip(table.rows, rows):
        keep_row_together(row)
        for idx, value in enumerate(values):
            if idx in (0, 2):
                set_cell_shading(row.cells[idx], LIGHT_GRAY)
                add_table_text(row.cells[idx], value, size=9.5, color=NAVY, bold=True)
            else:
                add_table_text(row.cells[idx], value, size=9.5, color=GRAY)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


def build_summary_table(doc):
    table = doc.add_table(rows=1, cols=4)
    set_table_geometry(table, [480, 1600, 6080, 1200])
    set_table_borders(table, BORDER, "6")
    header = table.rows[0]
    mark_header_row(header)
    for cell, value in zip(header.cells, ["#", "Area", "Requested change", "Priority"]):
        set_cell_shading(cell, LIGHT_BLUE)
        add_table_text(cell, value, size=9.5, color=NAVY, bold=True)
    for values in SUMMARY_ROWS:
        row = table.add_row()
        keep_row_together(row)
        for idx, value in enumerate(values):
            if idx == 0:
                add_table_text(row.cells[idx], value, size=9.5, color=NAVY, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
            elif idx == 1:
                add_table_text(row.cells[idx], value, size=9.5, color=NAVY, bold=True)
            elif idx == 3:
                color = "9B1C1C" if value == "High" else "7A5A00"
                add_table_text(row.cells[idx], value, size=9.5, color=color, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
            else:
                add_table_text(row.cells[idx], value, size=9.5, color=GRAY)
    # Reapply exact cell widths after adding body rows so every tcW matches the grid.
    set_table_geometry(table, [480, 1600, 6080, 1200])
    return table


def add_video_section(doc, video, num_id):
    add_heading(doc, f"{video['number']}. {video['title']}", level=2)
    add_label_paragraph(doc, "Source file: ", f"{video['file']}  |  Duration: {video['duration']}", size=9.5)
    add_heading(doc, "Arabic transcription", level=3)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.15
    set_bidi(p, True)
    r = p.add_run(video["arabic"])
    set_run_font(r, name="Arial", size=10.5, color="222222")
    add_heading(doc, "English translation", level=3)
    p = doc.add_paragraph(video["translation"])
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.10
    for r in p.runs:
        set_run_font(r, size=10.5, color="222222")
    add_heading(doc, "Client request captured", level=3)
    for item in video["changes"]:
        add_bullet(doc, item, num_id, size=10.5)


def build_doc():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    set_style_font(styles["Normal"], size=11, color="000000")
    styles["Normal"].paragraph_format.space_before = Pt(0)
    styles["Normal"].paragraph_format.space_after = Pt(6)
    styles["Normal"].paragraph_format.line_spacing = 1.10
    set_style_font(styles["Title"], size=25, color=NAVY, bold=True)
    set_style_font(styles["Subtitle"], size=13.5, color=GRAY)
    set_style_font(styles["Heading 1"], size=16, color=BLUE, bold=True)
    styles["Heading 1"].paragraph_format.space_before = Pt(16)
    styles["Heading 1"].paragraph_format.space_after = Pt(8)
    styles["Heading 1"].paragraph_format.keep_with_next = True
    set_style_font(styles["Heading 2"], size=13, color=BLUE, bold=True)
    styles["Heading 2"].paragraph_format.space_before = Pt(12)
    styles["Heading 2"].paragraph_format.space_after = Pt(6)
    styles["Heading 2"].paragraph_format.keep_with_next = True
    set_style_font(styles["Heading 3"], size=12, color=DARK_BLUE, bold=True)
    styles["Heading 3"].paragraph_format.space_before = Pt(8)
    styles["Heading 3"].paragraph_format.space_after = Pt(4)
    styles["Heading 3"].paragraph_format.keep_with_next = True

    doc.core_properties.title = "Client Feedback Videos - Transcription, Translation, and Requested Changes"
    doc.core_properties.subject = "Arabic client feedback translated into English with implementation requests"
    doc.core_properties.author = ""
    doc.core_properties.comments = ""

    # Running header and footer.
    header = section.header
    hp = header.paragraphs[0]
    hp.text = ""
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    hp.paragraph_format.space_after = Pt(0)
    r = hp.add_run("Client Feedback Review")
    set_run_font(r, size=9, color=MUTED, bold=True)
    r = hp.add_run("   |   Car Sales Platform")
    set_run_font(r, size=9, color=MUTED)
    footer = section.footer
    fp = footer.paragraphs[0]
    fp.text = ""
    fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    fp.paragraph_format.space_before = Pt(0)
    r = fp.add_run("Client Feedback Review  |  Page ")
    set_run_font(r, size=9, color=MUTED)
    add_page_field(fp, "PAGE")
    r = fp.add_run(" of ")
    set_run_font(r, size=9, color=MUTED)
    add_page_field(fp, "NUMPAGES")

    num_id = add_numbering_definition(doc)

    kicker = doc.add_paragraph()
    kicker.paragraph_format.space_before = Pt(10)
    kicker.paragraph_format.space_after = Pt(2)
    r = kicker.add_run("CLIENT FEEDBACK REVIEW")
    set_run_font(r, size=9.5, color=BLUE, bold=True)

    title = doc.add_paragraph()
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(5)
    title.paragraph_format.keep_with_next = True
    r = title.add_run("Arabic Feedback Videos")
    set_run_font(r, size=25, color=NAVY, bold=True)

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_before = Pt(0)
    subtitle.paragraph_format.space_after = Pt(16)
    subtitle.paragraph_format.keep_with_next = True
    r = subtitle.add_run("Transcription, English translation, and implementation-ready requested changes")
    set_run_font(r, size=13.5, color=GRAY)

    build_metadata_table(doc)

    note = doc.add_paragraph()
    note.paragraph_format.space_before = Pt(8)
    note.paragraph_format.space_after = Pt(12)
    note.paragraph_format.left_indent = Inches(0.08)
    note.paragraph_format.right_indent = Inches(0.08)
    note.paragraph_format.line_spacing = 1.10
    set_paragraph_shading(note, CALLOUT)
    set_paragraph_border(note, "left", BORDER, "18", "8")
    r = note.add_run("Review note: ")
    set_run_font(r, size=9.5, color=NAVY, bold=True)
    r = note.add_run(
        "The six clips are in Arabic with Saudi/Gulf dialect phrasing. The Arabic transcription has been lightly cleaned for punctuation and obvious speech-recognition slips. Text in square brackets marks a word that was not clear in the audio; translations preserve the intended product meaning."
    )
    set_run_font(r, size=9.5, color=GRAY)

    add_heading(doc, "Requested changes at a glance", level=1)
    p = doc.add_paragraph("The table below consolidates the feedback into implementation themes. Priorities are recommended based on usability impact and the number of videos that reference the issue.")
    p.paragraph_format.space_after = Pt(8)
    for r in p.runs:
        set_run_font(r, size=10.5, color=GRAY)
    build_summary_table(doc)

    add_heading(doc, "Implementation notes", level=2)
    for item in [
        "The color feedback appears in two clips and is grouped as one requirement: common-color selection plus support for custom and combined colors.",
        "The pricing clip is phrased as a question, so the requested action is to clarify the field behavior before finalizing the pricing flow.",
        "The inventory feedback asks for a complete post-creation lifecycle: edit, add missing details, and delete."
    ]:
        add_bullet(doc, item, num_id, size=10.5)

    # Start the detailed evidence on a fresh page.
    doc.add_page_break()
    add_heading(doc, "Video-by-video transcription and translation", level=1)
    p = doc.add_paragraph("Each section preserves the source file name, a cleaned Arabic transcript, the English translation, and the client request captured from that clip.")
    p.paragraph_format.space_after = Pt(8)
    for r in p.runs:
        set_run_font(r, size=10.5, color=GRAY)

    for video in VIDEOS:
        add_video_section(doc, video, num_id)

    add_heading(doc, "Recommended next step", level=1)
    p = doc.add_paragraph(
        "Convert the six consolidated themes into acceptance criteria for the vehicle search, dealer form fields, color selection, pricing flow, and inventory actions. Confirm the pricing-field interpretation with the client before development begins."
    )
    p.paragraph_format.space_after = Pt(6)
    for r in p.runs:
        set_run_font(r, size=11, color="222222")

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(f"Wrote {OUTPUT}")


if __name__ == "__main__":
    build_doc()
